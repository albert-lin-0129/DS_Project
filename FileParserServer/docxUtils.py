import os
from docx import Document
import utilsInterface
import Para_table_image_extraction

pwd = os.path.dirname(__file__)
UPLOAD_FOLDER = os.path.join(pwd, 'WPWPOI/files')
sort_array = []


class DocxUtils(utilsInterface.UtilsInterface):

    @staticmethod
    def get_table_paragraphs(file, tid):
        paragraphs = []
        table = file.tables[tid]
        for i in range(0, len(table.rows)):
            for j in range(0, len(table.columns)):
                for paragraph in table.cell(i, j).paragraphs:
                    if paragraph.text != "":
                        paragraphs.append(paragraph)
        return paragraphs

    @staticmethod
    def get_paragraphs(file):
        paragraphs = []
        global sort_array
        for p in file.paragraphs:
            if p.text != "":
                paragraphs.append(p)
        offset = 0
        for i in range(0, len(sort_array)):
            if sort_array[i][1] != 'Novalue':
                t_p = DocxUtils.get_table_paragraphs(file, int(sort_array[i][1]))
                sort_array[i][3] = len(t_p)
                for p in t_p:
                    paragraphs.insert(i + offset, p)
                    offset += 1

        return paragraphs

    @staticmethod
    def get_titles(file):
        titles = []
        font_sizes = []
        context_size = 2147000000
        for p in DocxUtils.get_paragraphs(file):
            runs = p.runs
            font_size = 0
            for run in runs:
                if run.font is not None and run.font.size is not None:
                    context_size = min(context_size, run.font.size)
                    font_size = max(font_size, run.font.size)
            font_sizes.append(font_size)

        index = 0
        for p in DocxUtils.get_paragraphs(file):
            if p.text != "" and font_sizes[index] > context_size:
                titles.append(index)
            index += 1
        return titles

    @staticmethod
    def get_tables(file):
        return file.tables

    @staticmethod
    def get_images(file):
        images = []
        inline_shapes = file.inline_shapes
        for shape in inline_shapes:
            image = {}
            image["height"] = shape.height
            image["width"] = shape.width
            image["type"] = shape.type
            images.append(image)
        # dict_rel = file.part._rels
        # for rel in dict_rel:
        #     rel = dict_rel[rel]
        #     if "image" in rel.target_ref:
        #         images.append(rel.target_part.image)
        return images

    @staticmethod
    def get_title_last_index(file, pid):
        paragraphs_info = DocxUtils.parse_all_paragraphs(file)
        titles_indexes = DocxUtils.get_titles(file)
        length = len(titles_indexes)
        last_index = len(paragraphs_info)
        for i in range(0, length):
            if titles_indexes[i] == pid:
                if i != length - 1:
                    last_index = titles_indexes[i + 1]
                else:
                    last_index = len(paragraphs_info)
        return last_index

    @staticmethod
    def parse_paragraphs_type_by_id(file, pid):
        """
        需要返回一个字典，里面key是格式特征，value是内容
        """
        res = {}
        paragraph_format = DocxUtils.get_paragraphs(file)[pid].paragraph_format
        res["paragraphAlignment"] = paragraph_format.alignment
        res["first_line_indent"] = paragraph_format.first_line_indent
        res["keep_together"] = paragraph_format.keep_together
        res["keep_with_next"] = paragraph_format.keep_with_next
        res["left_indent"] = paragraph_format.left_indent
        res["line_spacing"] = paragraph_format.line_spacing
        if paragraph_format.line_spacing is not None:
            res["line_spacing"] = paragraph_format.line_spacing
        return res

    @staticmethod
    def parse_paragraph_fonts_type_by_id(file, pid):
        """
        需要返回一个字典，里面key是字体信息，value是内容
        """
        res = {}
        runs = DocxUtils.get_paragraphs(file)[pid].runs
        for run in runs:
            if run.font is not None:
                res["paragraphFont"] = run.font.name
                res["fontSize"] = run.font.size
                if run.font.size is not None:
                    res["fontSize"] = run.font.size
                res["fontItalic"] = run.font.italic
                res["fontColor"] = run.font.color.rgb
                res["fontBold"] = run.font.bold
        return res

    @staticmethod
    def parse_paragraph_by_id(file, pid):
        """
        需要返回一个字典，里面key是特征，value是内容
        """
        res = {}
        paragraph = DocxUtils.get_paragraphs(file)[pid]
        res["paragraphId"] = pid
        res["paragraphText"] = paragraph.text

        dic = DocxUtils.parse_paragraphs_type_by_id(file, pid)
        for key in dic.keys():
            res[key] = dic[key]
        dic = DocxUtils.parse_paragraph_fonts_type_by_id(file, pid)
        for key in dic.keys():
            res[key] = dic[key]

        return res

    @staticmethod
    def parse_tables_by_table_id(file, tid):
        """
        需要返回一个字典，里面key是特征，value是内容
        """
        res = {}
        table = DocxUtils.get_tables(file)[tid]
        res["tableWidth"] = len(table.columns)
        res["tableHeight"] = len(table.rows)
        res["text"] = []
        table_p = []
        for i in range(0, len(table.rows)):
            row = []
            for j in range(0, len(table.columns)):
                row.append(table.cell(i, j).text)
                for paragraph in table.cell(i, j).paragraphs:
                    if paragraph.text != "":
                        table_p.append(paragraph)
            res["text"].append(row)

        pid_list = []
        has_id = []
        file_p = DocxUtils.get_paragraphs(file)
        length = len(file_p)
        for i in range(0, length):
            for p in table_p:
                if file_p[i] == p and not has_id.__contains__(p):
                    pid_list.append(i)
                    has_id.append(p)
        res["paragraphIds"] = pid_list
        return res

    @staticmethod
    def parse_all_paragraphs(file):
        """
        需要返回一个字典类型的数组，每个字典是一个段落内容
        """
        res = []
        paragraphs = DocxUtils.get_paragraphs(file)
        for index in range(0, len(paragraphs)):
            res.append(DocxUtils.parse_paragraph_by_id(file, index))
        return res

    @staticmethod
    def parse_all_tables(file):
        """
        需要返回一个字典类型的数组，每个字典是一个表格内容
        """
        res = []
        tables = DocxUtils.get_tables(file)
        for index in range(0, len(tables)):
            res.append(DocxUtils.parse_tables_by_table_id(file, index))
        return res

    @staticmethod
    def parse_all_pics(file):
        """
        需要返回一个字典类型的数组，每个字典是一个图片内容
        """
        # res = []
        # for index in range(0, len(file.tables)):
        #     res.append(DocxUtils.parse_tables_by_pics_id(file, index))
        res = DocxUtils.get_images(file)
        return res

    @staticmethod
    def parse_all_titles(file):
        """
        需要返回一个字典类型的数组，每个字典是一个标题内容
        """
        res = []
        titles = DocxUtils.get_titles(file)
        for pid in titles:
            res.append(DocxUtils.parse_paragraph_by_id(file, pid))
        return res

    @staticmethod
    def get_file(path):
        """
        需要返回一个file，就是解析后的一个类，之后传入每一个方法中
        docx就是Document类
        """
        file = Document(path)
        Para_table_image_extraction.document = file
        global sort_array
        sort_array = Para_table_image_extraction.get_combine_dataframe()
        return file

    @staticmethod
    def parse_paragraph_by_title_id(file, pid):
        """
        需要返回一个字典类型的数组，里面key是标题特征，value是内容
        """
        res = []
        paragraphs_info = DocxUtils.parse_all_paragraphs(file)
        last_index = DocxUtils.get_title_last_index(file, pid)
        is_in_title = False
        for info in paragraphs_info:
            if info["paragraphId"] == pid:
                is_in_title = True
            if info["paragraphId"] == last_index:
                break
            if is_in_title:
                res.append(info)
        return res

    @staticmethod
    def parse_paragraph_pics_by_title_id(file, pid):
        """
        需要返回一个字典类型的数组，里面key是图片特征，value是内容
        """
        images = DocxUtils.get_images(file)
        last_index = DocxUtils.get_title_last_index(file, pid)
        res = []
        is_in_title = False
        index = 0
        image_index = 0
        for info in sort_array:
            if is_in_title:
                if sort_array[0].lower().__contains__("image"):
                    res.append(images[image_index])
            if sort_array[0].lower().__contains__("image"):
                image_index += 1
                continue
            if info[3] >= index:
                is_in_title = True
            if info[3] >= last_index:
                break

        return res
        pass

    @staticmethod
    def parse_paragraph_tables_by_title_id(file, pid):
        """
        需要返回一个字典类型的数组，里面key是表格特征，value是内容
        """

        tables = DocxUtils.parse_all_tables(file)
        last_index = DocxUtils.get_title_last_index(file, pid)
        res = []
        for t in tables:
            for t_pid in t["paragraphIds"]:
                if t_pid in range(pid, last_index):
                    res.append(t)
        return res


if __name__ == '__main__':
    filePath = os.path.join(UPLOAD_FOLDER, "test.docx")
    # print(DocxUtils.parse_all_titles(DocxUtils.get_file(filePath)))
    print(DocxUtils.parse_paragraph_by_title_id(DocxUtils.get_file(filePath), 3))
    # aimages = DocxUtils.get_images(DocxUtils.get_file(filePath))
    # atable = DocxUtils.parse_tables_by_table_id(DocxUtils.get_file(filePath), 0)
    # print(DocxUtils.parse_tables_by_table_id(DocxUtils.get_file(filePath), 0))
    #
    # for image in aimages:
    #     print(image["height"])
    # print(DocxUtils.parse_all_paragraphs(DocxUtils.get_file(filePath)))
    # print(DocxUtils.parse_paragraph_by_id(DocxUtils.get_file(filePath), 0))
    # print(combined_df, '\n')

    # docx = Document(filePath)

    # doc_pars = DocxUtils.get_paragraphs(docx)
    # print(doc_pars)
    # DocxUtils.parse_images(docx)
    # DocxUtils.parse_paragraph_type(doc_pars[1])
