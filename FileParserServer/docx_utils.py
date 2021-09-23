import os
from docx import Document

pwd = os.path.dirname(__file__)
UPLOAD_FOLDER = os.path.join(pwd, 'WPWPOI/files')


class DocxUtils:

    @staticmethod
    def get_paragraphs(document):
        paragraphs = {}
        i = 0
        for paragraph in document.paragraphs:
            # print(paragraph.text)
            paragraphs[i] = paragraph
            i += 1
        return paragraphs

    @staticmethod
    def get_tables(document):
        tables = {}
        i = 0
        for table in document.tables:
            # print(paragraph.text)
            tables[i] = table
            i += 1
        return tables

    @staticmethod
    def parse_paragraph_type(paragraph):
        print(paragraph.text)
        print(paragraph.paragraph_format.alignment)
        print(paragraph.paragraph_format.first_line_indent)
        print(paragraph.style.font.name)

    @staticmethod
    def parse_images(document):
        """
        image需要知道这张图片的段落位置以及
        """
        dict_rel = document.part._rels
        for rel in dict_rel:
            rel = dict_rel[rel]
            if "image" in rel.target_ref:
                print(rel.target_part)


if __name__ == '__main__':
    path = os.path.join(UPLOAD_FOLDER, "test.docx")
    docx = Document(path)

    doc_pars = DocxUtils.get_paragraphs(docx)
    print(doc_pars)
    DocxUtils.parse_images(docx)
    DocxUtils.parse_paragraph_type(doc_pars[1])
